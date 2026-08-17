"""docx 使用说明书生成（python-docx 可选依赖，懒加载）。

对齐旧版 generate_docx_manual：封面 + 目录 + 六章正文（核心功能/智能预测/系统管理/FAQ/技术支持）。
python-docx 未安装时抛 RuntimeError（UI 降级提示）。
"""

from pathlib import Path
from typing import Optional

from smart_farm.services.instruction_data import FAQ, SUPPORT, get_full_instruction


def generate_docx_manual(output_path: Optional[Path] = None) -> str:
    """生成使用说明书 docx，返回输出路径。"""
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "未安装 python-docx。请运行 `uv pip install python-docx` 后重试。"
        ) from exc

    output_path = output_path or Path(__file__).resolve().parents[3] / "docs" / "使用说明书.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    # 页面边距
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.5)
        section.top_margin = section.bottom_margin = Inches(1.5)

    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("智慧大棚数据管理平台")
    run.bold = True
    run.font.size = Pt(24)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run("使用说明书")
    srun.font.size = Pt(20)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("版本：V1.0\n日期：2026 年")
    doc.add_page_break()

    # 目录（手动章节列表）
    doc.add_heading("目录", level=1)
    toc_entries = ["一、核心功能", "二、智能分析与预测", "三、系统管理", "四、常见问题", "五、技术支持"]
    for entry in toc_entries:
        doc.add_paragraph(entry)

    # 章节内容
    instruction = get_full_instruction()
    chapter_names = {"核心功能": "一、核心功能", "智能分析与预测": "二、智能分析与预测", "系统管理": "三、系统管理"}
    for category, display in chapter_names.items():
        doc.add_heading(display, level=1)
        for feature in instruction.get(category, []):
            doc.add_heading(feature["name"], level=2)
            doc.add_paragraph(f"【描述】{feature['desc']}")
            doc.add_paragraph("【操作步骤】")
            for step in feature["steps"]:
                doc.add_paragraph(f"- {step}", style="List Bullet")
            if feature["notes"]:
                doc.add_paragraph(f"【注意事项】{feature['notes']}")
        doc.add_page_break()

    doc.add_heading("四、常见问题", level=1)
    for q, a in FAQ:
        doc.add_paragraph(f"Q: {q}")
        doc.add_paragraph(f"A: {a}")

    doc.add_heading("五、技术支持", level=1)
    doc.add_paragraph(f"邮箱：{SUPPORT['email']}")
    doc.add_paragraph(f"项目：{SUPPORT['repo']}")

    doc.save(str(output_path))
    return str(output_path)
